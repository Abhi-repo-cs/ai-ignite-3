import streamlit as st
import pandas as pd
import time
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import json
import os
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4

from groq import Groq
client = Groq(api_key=st.secrets["GROQ_API_KEY"])




st.set_page_config(
    page_title="AI Powered Question Bank Generator",
    page_icon="📘",
    layout="wide"
)


st.markdown("""
<style>

html, body, [class*="css"]  {
    font-family: 'Segoe UI', sans-serif;
}

.card {
    background: rgba(255,255,255,0.75);
    border-radius: 15px;
    padding: 25px;
    box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    margin-bottom: 20px;
}

h1 {
    text-shadow: 0px 0px 18px rgba(255,75,43,0.3);
}

[data-testid="stSidebar"] {
    background-color: #0B0F19;
    color: white;
}

div.stButton > button:first-child {
    background: linear-gradient(90deg,#FF416C,#FF4B2B);
    color:white;
    border-radius:10px;
    padding:12px 22px;
    font-weight:700;
    border:0px;
    box-shadow:0px 0px 12px rgba(255,75,43,0.5);
}
div.stButton > button:first-child:hover {
    transform:scale(1.03);
}

</style>
""", unsafe_allow_html=True)


st.markdown("""
<style>

[data-testid="stSidebar"] * {
    color: white;
}

div[role="radiogroup"] label p {
    font-size: 16px;
    font-weight: 600;
}

div[role="radiogroup"] > label:hover {
    background-color: rgba(255,255,255,0.08);
    border-radius: 8px;
    padding: 4px;
}

</style>
""", unsafe_allow_html=True)



def detect_bloom(question):
    prompt = f"""
Classify the Bloom level for this question only as one word:
Remember, Understand, Apply, Analyze, Evaluate, Create

Question:
{question}
"""

    r = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[{"role":"user","content":prompt}]
    )

    return r.choices[0].message.content.strip()


def check_duplicate(question, existing_questions):

    if len(existing_questions) == 0:
        return 0

    docs = existing_questions + [question]

    vec = TfidfVectorizer().fit_transform(docs)
    sim_matrix = cosine_similarity(vec)

    similarity_scores = sim_matrix[-1][:-1]

    return max(similarity_scores)

#PLANNING AGENT
def generate_plan(syllabus, cos, q_count, bloom, difficulty, marks):

    prompt = f"""
You are an Assessment Design Expert.
Create a QUESTION PLAN before generating questions.

Inputs:
Syllabus: {syllabus}
Course Outcomes: {cos}

Requirements:
Total Questions: {q_count}
Target Bloom Level: {bloom}
Difficulty: {difficulty}
Marks per Question: {marks}

Output the plan in clear bullet points including:

1. Intended Question Coverage Strategy
2. Bloom Level Mix Justification
3. Difficulty Spread Strategy
4. CO Mapping Strategy
5. Risk Controls for Duplicates & Relevance
"""

    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return r.choices[0].message.content

#AUDIT AGENT
def audit_question(question, syllabus, cos):

    prompt = f"""
You are an Assessment Quality Auditor.

Evaluate the following exam question:

Question:
{question}

Context:
Syllabus: {syllabus}
Course Outcomes: {cos}

Check and respond in plain text with:

1. Relevance (High / Medium / Low)
2. Bloom Level (1 word)
3. Clarity (Clear / Needs Improvement)
4. Difficulty (Easy / Medium / Hard)
5. Suggested Improvement (if any)

Keep it short.
"""

    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return r.choices[0].message.content

#REPAI RAGENT
def repair_question(question, audit_feedback, syllabus, cos):

    prompt = f"""
You are an AI Assessment Improver.

Original Question:
{question}

Audit Feedback:
{audit_feedback}

Syllabus:
{syllabus}

Course Outcomes:
{cos}

TASK:
Rewrite the question to FIX all weaknesses
while keeping meaning and marks level SAME.

Return ONLY the improved question.
"""

    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}]
    )

    return r.choices[0].message.content.strip()



def save_data():
    with open("questions.json","w") as f:
        json.dump(st.session_state.questions,f)


def load_data():
    if os.path.exists("questions.json"):
        with open("questions.json","r") as f:
            st.session_state.questions = json.load(f)


def export_to_docx(data):
    doc = Document()
    doc.add_heading('Question Bank', 0)

    for i, q in enumerate(data):
        doc.add_paragraph(f"Q{i+1}. {q['Question']}")
        doc.add_paragraph(f"Bloom Level: {q['Bloom']}")
        doc.add_paragraph(f"Difficulty: {q['Difficulty']}")
        doc.add_paragraph(f"Marks: {q['Marks']}")
        doc.add_paragraph(" ")

    file_path = "question_bank.docx"
    doc.save(file_path)
    return file_path


def export_to_pdf(data):
    file_path = "question_bank.pdf"
    c = canvas.Canvas(file_path, pagesize=A4)
    width, height = A4

    y = height - 40

    c.setFont("Helvetica", 11)
    c.drawString(50, y, "Question Bank")
    y -= 30

    for i, q in enumerate(data):
        lines = [
            f"Q{i+1}. {q['Question']}",
            f"Bloom Level: {q['Bloom']}",
            f"Difficulty: {q['Difficulty']}",
            f"Marks: {q['Marks']}",
            ""
        ]

        for line in lines:
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - 50
            c.drawString(50, y, line)
            y -= 20

    c.save()
    return file_path



st.sidebar.image("static/logo.webp", width=120)

st.sidebar.title(" Navigation")

page = st.sidebar.radio(
    "",
    ["Home","Generate","Question Bank","Audit Dashboard","About Us","Rate Us"]
)

st.sidebar.write("---")

st.sidebar.subheader(" Team Members")
st.sidebar.write("""
• Abhishek  
• Shahid Asmar  
• Muralidharan  
• Hariharan
""")


if "questions" not in st.session_state:
    st.session_state.questions = []

load_data()

if "ratings" not in st.session_state:
    st.session_state.ratings = []


if page == "Home":

    st.markdown(
        """
        <div style="text-align:center;">
            <h1> Question Bank Generator</h1>
            <p style="font-size:20px;">
                <i>Outcome-Aligned • Efficient • Futuristic AI Assistant</i>
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.write("")
    st.write("")

    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("🙏 Vanakkam")
    st.write("""
AI Question Bank Generator is built to support educators in designing fair,
structured and outcome-oriented assessments effortlessly. Instead of manually
drafting questions and mapping them to outcomes, our tool helps you generate,
organize and audit question banks in just a few clicks.
    """)
    st.markdown("</div>", unsafe_allow_html=True)


    st.markdown("<div class='card'>", unsafe_allow_html=True)
    st.subheader("✨ Key Features")
    st.write("""
- 🎯 Outcome aligned  
- 🧠 Bloom Taxonomy classified  
- ⚖ Exam-Ready balance  
- 📊 Analytics & frequencies  
    """)
    st.markdown("</div>", unsafe_allow_html=True)

    st.success("Use the left navigation to be amazed ")


#GENERATE
elif page == "Generate":

    st.header(" Generate Questions")

    syllabus = st.text_area("📘 Paste Your Syllabus", height=120)

    cos = st.text_area(" Course Outcomes (one per line)", height=120)

    bloom = st.selectbox(
        " Target Bloom Level",
        ["Remember","Understand","Apply","Analyze","Evaluate","Create"]
    )

    q_count = st.slider(" Questions to Generate", 1, 20, 5)

    difficulty = st.radio(" Difficulty", ["Easy","Medium","Hard","Mixed"])

    marks = st.selectbox(" Marks per Question", [2, 8, 11])

    st.write("---")

    status = st.empty()

    if st.button(" GENERATE", key="gen_btn"):

        with st.spinner("️ Generating questions… please wait"):
            plan = generate_plan(syllabus, cos, q_count, bloom, difficulty, marks)

            st.subheader("AI Strategy Plan")
            st.write(plan)

            st.write("---")

            time.sleep(1.5)

            prompt = f"""
Generate {q_count} exam questions for the syllabus below.

Syllabus:
{syllabus}

Course Outcomes:
{cos}

Bloom Level Target: {bloom}
Difficulty: {difficulty}
Marks: {marks}

Return questions in numbered list.
"""

            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[{"role":"user","content":prompt}]
            )
            
            text = response.choices[0].message.content
            questions_text = text.split("\n")


            for q in questions_text:
                if q.strip():
                    existing = [x["Question"] for x in st.session_state.questions]

                    dup_score = check_duplicate(q, existing)

                    flag = "⚠ Possible Duplicate" if dup_score > 0.80 else " Unique"

                    bloom_level = detect_bloom(q)
                    audit_result = audit_question(q, syllabus, cos)

                    if "Needs Improvement" in audit_result or "Low" in audit_result:
                        improved_question = repair_question(q, audit_result, syllabus, cos)

                        final_question = improved_question
                        repair_status = "Improved by AI"

                    else:
                        final_question = q
                        repair_status = "✔ Accepted"


                    st.session_state.questions.append({
                        "Question": final_question.strip(),
                        "Bloom": bloom_level,
                        "Difficulty": difficulty,
                        "Marks": marks,
                        "Duplicate Risk": dup_score,
                        "Flag": flag,
                        "Audit Feedback": audit_result,
                        "Status": repair_status})


            save_data()

            status.success(" Questions added to bank successfully!")



elif page == "Question Bank":

    st.header(" Your Question Bank")

    if len(st.session_state.questions)==0:
        st.warning("No questions yet. Try generating some 🙂")
    else:
        df = pd.DataFrame(st.session_state.questions)
        st.dataframe(df, use_container_width=True)


        st.download_button(
            "⬇️ Download as CSV",
            df.to_csv(index=False),
            "question_bank.csv"
        )

        doc_file = export_to_docx(st.session_state.questions)
        pdf_file = export_to_pdf(st.session_state.questions)

        st.download_button(
            "📄 Download as Word (DOCX)",
            data=open(doc_file, "rb"),
            file_name="question_bank.docx",
        )

        st.download_button(
            "🧾 Download as PDF",
            data=open(pdf_file, "rb"),
            file_name="question_bank.pdf",
        )



elif page == "Audit Dashboard":

    st.header("📊 Previous Question Paper Analytics")

    if len(st.session_state.questions)==0:
        st.warning("No data available yet.")
    else:
        df = pd.DataFrame(st.session_state.questions)

        st.subheader("🧠 Bloom Distribution")
        st.bar_chart(df["Bloom"].value_counts())

        st.subheader("⚖ Difficulty Distribution")
        st.bar_chart(df["Difficulty"].value_counts())

        st.subheader("🚨 Duplicate Risk Levels")
        st.bar_chart(df["Duplicate Risk"])

        st.subheader("📎 Summary Metrics")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.metric("Total Questions", len(df))

        with col2:
            st.metric("Duplicate Alerts", sum(df["Flag"]=="⚠ Possible Duplicate"))

        with col3:
            st.metric("Average Duplicate Risk", round(df["Duplicate Risk"].mean(),2))



elif page == "About Us":

    st.header("️ About This Project")

    st.write("""
This project was created as part of IgniteHack 2026  
with the goal of empowering educators using intelligent assessment tools.
""")

    st.subheader(" Team")
    st.write("""
• Abhishek  
• Shahid Asmar  
• Muralidharan  
• Hariharan
""")

    st.write("---")

    st.caption(" Built with ❤️ and Streamlit")



elif page == "Rate Us":

    st.header("⭐ Rate Our Application")

    st.write("""
We hope this tool has been helpful 🙂
Please rate your experience — your feedback helps us improve.
""")

    rating = st.slider("Your Rating", 1, 5, 4)

    comments = st.text_area("💬 Additional Comments ")

    if st.button("Submit Feedback", key="feedback_btn"):

        st.session_state.ratings.append({
            "rating": rating,
            "comment": comments
        })

        st.success("✅ Thank you for your feedback! 🙏")

    st.write("---")

    if len(st.session_state.ratings) > 0:

        avg = sum(r["rating"] for r in st.session_state.ratings) / len(st.session_state.ratings)

        st.subheader("📊 Rating Summary")

        st.metric("Average Rating", f"{avg:.2f} ⭐")

        st.write(f"Total Reviews: {len(st.session_state.ratings)}")

        st.write("---")

        st.subheader("📝 Recent Feedback")

        for r in reversed(st.session_state.ratings[-5:]):
            st.write(f"⭐ {r['rating']}")
            st.write(r["comment"])
            st.write("---")




st.markdown("""
<hr>
<div style='text-align:center'>
    <p> AI Powered Question Bank Generator — IgniteHack 2026</p>
</div>
""", unsafe_allow_html=True)
